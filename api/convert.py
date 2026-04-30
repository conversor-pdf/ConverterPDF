from flask import Flask, request, send_file, jsonify
import fitz  # PyMuPDF
import io
import zipfile
import os
import sys
from PIL import Image
from docx import Document
from docx.shared import Inches, Cm

app = Flask(__name__)

@app.route('/api/convert', methods=['POST'])
@app.route('/', methods=['POST'])
def convert():
    print("Iniciando conversão...", file=sys.stderr)
    try:
        if 'files' not in request.files:
            print("Erro: Nenhum arquivo na requisição", file=sys.stderr)
            return jsonify({"success": False, "error": "Nenhum arquivo enviado"}), 400
        
        files = request.files.getlist('files')
        target_format = request.form.get('format', 'png')
        print(f"Arquivos recebidos: {len(files)}, Formato: {target_format}", file=sys.stderr)
        
        output_files = []
        
        for file in files:
            if not file.filename:
                continue
            
            print(f"Processando arquivo: {file.filename}", file=sys.stderr)
            pdf_data = file.read()
            
            if not pdf_data:
                print(f"Aviso: Arquivo {file.filename} está vazio", file=sys.stderr)
                continue
                
            try:
                doc = fitz.open(stream=pdf_data, filetype="pdf")
                
                for i in range(len(doc)):
                    page = doc.load_page(i)
                    pix = page.get_pixmap()
                    img_data = pix.tobytes(target_format)
                    output_files.append({
                        "name": f"{os.path.splitext(file.filename)[0]}_pag_{i+1}.{target_format}",
                        "data": img_data
                    })
                doc.close()
            except Exception as pdf_err:
                print(f"Erro ao abrir PDF {file.filename}: {str(pdf_err)}", file=sys.stderr)
                return jsonify({"success": False, "error": f"Erro no arquivo {file.filename}: {str(pdf_err)}"}), 400

        if not output_files:
            print("Erro: Nenhuma imagem gerada", file=sys.stderr)
            return jsonify({"success": False, "error": "Nenhuma imagem gerada. Verifique se os arquivos são PDFs válidos."}), 400

        print(f"Sucesso! Geradas {len(output_files)} imagens", file=sys.stderr)
        if len(output_files) == 1:
            return send_file(
                io.BytesIO(output_files[0]["data"]),
                mimetype=f'image/{target_format}',
                as_attachment=True,
                download_name=output_files[0]["name"]
            )
        
        # Criar ZIP para múltiplos arquivos/páginas
        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w') as zf:
            for f in output_files:
                zf.writestr(f["name"], f["data"])
        
        memory_file.seek(0)
        return send_file(
            memory_file,
            mimetype='application/zip',
            as_attachment=True,
            download_name='conversoes_wtech.zip'
        )
    except Exception as e:
        print(f"Erro fatal: {str(e)}", file=sys.stderr)
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/export-grid', methods=['POST'])
def export_grid():
    try:
        export_format = request.form.get('format', 'pdf').lower()
        image_files = request.files.getlist('images')
        
        # A4 size at 150 DPI (approx 1240 x 1754 px)
        a4_w, a4_h = 1240, 1754
        cols, rows = 5, 3
        cell_w = a4_w // cols
        cell_h = a4_h // rows
        
        # Prepare the canvas
        canvas = Image.new('RGB', (a4_w, a4_h), 'white')
        
        for i, img_file in enumerate(image_files[:15]):
            img = Image.open(img_file)
            # Resize keeping proportions or fill? User said "preencher".
            # Let's use resize to fill the rectangle.
            img_resized = img.resize((cell_w, cell_h), Image.Resampling.LANCZOS)
            
            x = (i % cols) * cell_w
            y = (i // cols) * cell_h
            canvas.paste(img_resized, (x, y))
            
        output = io.BytesIO()
        
        if export_format == 'pdf':
            canvas.save(output, format='PDF')
            mimetype = 'application/pdf'
            filename = 'gabarito_wtech.pdf'
        elif export_format == 'docx':
            doc = Document()
            # Set margins to 0 for maximum space? No, let's keep some.
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            
            table = doc.add_table(rows=rows, cols=cols)
            
            for i, img_file in enumerate(image_files[:15]):
                img_file.seek(0)
                row = i // cols
                col = i % cols
                cell = table.cell(row, col)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                
                # Temp buffer for this image
                img_temp = io.BytesIO(img_file.read())
                # Scale width to fit column (approx 4cm)
                run.add_picture(img_temp, width=Cm(4))
                
            doc.save(output)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = 'gabarito_wtech.docx'
        elif export_format in ['png', 'jpg', 'jpeg']:
            fmt = 'JPEG' if export_format in ['jpg', 'jpeg'] else 'PNG'
            canvas.save(output, format=fmt)
            mimetype = f'image/{export_format}'
            filename = f'gabarito_wtech.{export_format}'
        else:
            return jsonify({"success": False, "error": "Formato não suportado"}), 400
            
        output.seek(0)
        return send_file(output, mimetype=mimetype, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Erro no export-grid: {str(e)}", file=sys.stderr)
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
